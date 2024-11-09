import pandas as pd
from docx.shared import Pt
from docx import Document
from datetime import datetime, timedelta
import logging
import os

def format_datetime_range(start_str, end_str, minutes_to_subtract):
    """Форматирует диапазон дат и времени в нужный формат."""
    date_time_format = "%d/%m/%Y %H:%M"
    original_start = datetime.strptime(start_str, date_time_format)
    original_end = datetime.strptime(end_str, date_time_format)

    adjusted_start = original_start - timedelta(minutes=minutes_to_subtract)

    start_time_str = adjusted_start.strftime("%H:%M")
    end_time_str = original_end.strftime("%H:%M")

    if adjusted_start.date() == original_end.date():
        date_str = adjusted_start.strftime("%d.%m.%Y")
        return f"с {start_time_str} до {end_time_str} {date_str}"
    else:
        start_date_str = adjusted_start.strftime("%d.%m.%Y")
        end_date_str = original_end.strftime("%d.%m.%Y")
        return f"с {start_time_str} {start_date_str} до {end_time_str} {end_date_str}"

def get_duration_string(total_minutes):
    """Возвращает строку длительности в часах и минутах с правильным склонением."""
    def get_hour_string(hours):
        if hours == 1:
            return "1 час"
        elif 2 <= hours <= 4:
            return f"{hours} часа"
        else:
            return f"{hours} часов"

    def get_minute_string(minutes):
        if minutes == 1:
            return "1 минута"
        elif 2 <= minutes <= 4:
            return f"{minutes} минуты"
        else:
            return f"{minutes} минут"

    hours = total_minutes // 60
    minutes = total_minutes % 60

    hours_str = get_hour_string(hours) if hours > 0 else ""
    minutes_str = get_minute_string(minutes) if minutes > 0 else ""

    return f"{hours_str} {minutes_str}".strip()

def add_paragraph_to_document(document, gas_name, duration_str, details, is_last=False):
    """Добавляет абзац в документ с заданными параметрами."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    ending = "." if is_last else ";"
    run = paragraph.add_run(f"по {gas_name} – {details}{ending}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

def record_max_excess_duration(path, file_name, document, gas_names, is_last=False):
    """Находит строки с максимальным количеством точек и записывает результаты в одну строку в документ."""
    try:
        df = pd.read_excel(f"{path}{file_name}.xlsx")
        
        # Добавляем определение региона
        mo_stations = ["МО", "Звенигород", "Балашиха-Салтыковка", "Реутов-2", "М (Балашиха-Речная)"]
        df["Регион"] = df.iloc[:, 0].apply(
            lambda x: "Московская область" if any(station in x for station in mo_stations) else "Москва"
        )

        details_parts = []
        
        for region in ["Москва", "Московская область"]:
            region_df = df[df["Регион"] == region]
            if not region_df.empty:
                max_points = region_df['Количество точек'].max()
                max_excess_rows = region_df[region_df['Количество точек'] == max_points].values.tolist()

                station_names = [line[0] for line in max_excess_rows]
                stations_str = ', '.join(station_names)

                first_line = max_excess_rows[0]
                duration_minutes = int(first_line[2]) * 20
                duration_str = get_duration_string(duration_minutes)
                formatted_range = format_datetime_range(first_line[3], first_line[4], 20)
                
                details_parts.append(f"{duration_str} {formatted_range} ({stations_str})")

        if details_parts:
            details = ", ".join(details_parts)
            add_paragraph_to_document(document, gas_names[file_name], details_parts[0], details, is_last)
            
    except FileNotFoundError:
        logging.warning(f"Файл {file_name}.xlsx не найден в директории {path}")
    except Exception as e:
        logging.error(f"Ошибка при обработке файла {file_name}.xlsx: {str(e)}")

def record_total_excess_duration(path, file_name, document, gas_names, is_last=False):
    """Находит суммарную длительность превышений для каждой точки и записывает результаты в одну строку в документ."""
    try:
        df = pd.read_excel(f"{path}{file_name}.xlsx")
        
        # Добавляем определение региона
        mo_stations = ["МО", "Звенигород", "Балашиха-Салтыковка", "Реутов-2", "М (Балашиха-Речная)"]
        df["Регион"] = df.iloc[:, 0].apply(
            lambda x: "Московская область" if any(station in x for station in mo_stations) else "Москва"
        )

        details_parts = []
        
        for region in ["Москва", "Московская область"]:
            region_df = df[df["Регион"] == region]
            if not region_df.empty:
                points_map = region_df.groupby(region_df.columns[0])['Количество точек'].sum().to_dict()
                max_value = max(points_map.values())

                max_keys = [key for key, value in points_map.items() if value == max_value]
                stations_str = ', '.join(max_keys)

                total_duration_minutes = max_value * 20
                duration_str = get_duration_string(total_duration_minutes)
                
                details_parts.append(f"{duration_str} ({stations_str})")

        if details_parts:
            details = ", ".join(details_parts)
            add_paragraph_to_document(document, gas_names[file_name], details_parts[0], details, is_last)

    except FileNotFoundError:
        logging.warning(f"Файл {file_name}.xlsx не найден в директории {path}")
    except Exception as e:
        logging.error(f"Ошибка при обработке файла {file_name}.xlsx: {str(e)}")

def add_custom_text(document, text, font_name='Times New Roman', font_size=14, bold=False):
    """Добавляет заголовок с указанным шрифтом, размером и жирным стилем."""
    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(0)
    heading.paragraph_format.space_before = Pt(0)
    run = heading.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold

def get_available_files(path, expected_files):
    """Возвращает список доступных файлов из ожидаемого списка."""
    available_files = []
    for file_name in expected_files:
        if os.path.exists(f"{path}{file_name}.xlsx"):
            available_files.append(file_name)
    return available_files

def process_multiple_files(path, document):
    """Обрабатывает доступные файлы и записывает результаты в один документ."""
    gas_names = {
        "CO_п": "оксиду углерода",
        "H2S_п": "сероводороду",
        "NO_п": "оксиду азота",
        "NO2_п": "диоксиду азота",
        "PM10_п": "PM₁₀",
        "C10H8_п": "нафталину",
        "C6H5OH_п": "фенолу",
        "C6H6_п": "бензолу",
        "C7H8_п": "толуолу",
        "C8H8_п": "стиролу", 
        "CH2O_п": "формальдегиду"
    }

    expected_files = list(gas_names.keys())
    available_files = get_available_files(path, expected_files)

    if not available_files:
        logging.warning("Не найдено ни одного Excel файла для обработки")
        return

    add_custom_text(document, 'Максимальная непрерывная длительность превышений:', 
                    font_name='Times New Roman', font_size=14, bold=True)
    for i, file_name in enumerate(available_files):
        is_last = (i == len(available_files) - 1)
        try:
            record_max_excess_duration(path, file_name, document, gas_names, is_last)
        except Exception as e:
            logging.error(f"Ошибка при обработке файла {file_name}: {str(e)}")
            continue

    # Максимальная общая длительность
    add_custom_text(document, 'Максимальная общая длительность превышений:', 
                    font_name='Times New Roman', font_size=14, bold=True)
    for i, file_name in enumerate(available_files):
        is_last = (i == len(available_files) - 1)
        try:
            record_total_excess_duration(path, file_name, document, gas_names, is_last)
        except Exception as e:
            logging.error(f"Ошибка при обработке файла {file_name}: {str(e)}")
            continue

def process_part2(directory_name, output_file='result.docx'):
    logging.info(f"Начало обработки part2 с входной директорией {directory_name}")
    document = Document()
    process_multiple_files(directory_name, document)
    
    document.save(output_file)
    logging.info(f"Документ Word сохранен: {output_file}")
    logging.info(f"Файл {output_file} создан: {os.path.exists(output_file)}")
    
    return output_file

def main():
    path = "C:/Users/HUAWEI/Desktop/eco_detector/data/part2/"
    output_file = process_part2(path)
    logging.info(f"Обработка part2 завершена. Результат сохранен в {output_file}")

if __name__ == "__main__":
    main()