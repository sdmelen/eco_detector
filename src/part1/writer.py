import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def create_word_document(formatted_results: str, output_file: str):
 """
 Создает документ Word на основе отформатированных результатов.

 :param formatted_results: Отформатированная строка с результатами
 :param output_file: Путь для сохранения документа Word
 """
 doc = Document()
 gas_names ={
    "CO": "оксиду углерода",
    "H2S": "сероводороду",
    "NO": "оксиду азота",
    "NO2": "диоксиду азота",
    "PM10": "PM10"
 }

 # Установка стилей
 styles = doc.styles
 style_normal = styles['Normal']
 style_normal.font.name = 'Times New Roman'
 style_normal.font.size = Pt(14)

 # Добавление заголовка
 title = doc.add_paragraph()
 title_run = title.add_run("Данные АСКЗА")
 title_run.bold = True
 title_run.font.size = Pt(14)
 title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

 for line in formatted_results.splitlines():
        p = doc.add_paragraph()
        if line == "Превышения ПДКмр:":  # Делаем заголовок жирным
            run = p.add_run(line)
            run.bold = True
        elif line in ("Москва", "МО"):
            run = p.add_run(line)
            run.italic = True
        else:
            # Проверяем, является ли строка подзаголовком с "по {газ} на"
            if line.startswith("по ") and " АСКЗА:" in line: 
                parts = line.split(" АСКЗА:")
                gas_short = parts[0].split('по ')[1] # находим краткое название газа
                gas_name = gas_names.get(gas_short, gas_short)  # заменяем его на полное
                run = p.add_run(f"по {gas_name} АСКЗА:")
                run.bold = True
                if len(parts) > 1:
                    run = p.add_run(parts[1])
            else:
                parts = line.split('на уровне 1 ПДКмр ')
                if len(parts) > 1:
                    run = p.add_run(parts[0])
                    run = p.add_run('на уровне 1 ПДКмр ')
                    run.bold = True
                    run = p.add_run(parts[1])
                else:
                    parts = line.split("до ")
                    if len(parts) > 1:
                        run = p.add_run(parts[0])
                        run = p.add_run(f"до {parts[1].split(' ПДКмр ')[0]} ПДКмр ")
                        run.bold = True
                        run = p.add_run(parts[1].split(' ПДКмр ')[1])
                    else:
                        run = p.add_run(line)

        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

 # Сохранение документа
 doc.save(output_file)
 logging.info(f"Документ Word сохранен: {output_file}")

def main(formatted_results: str, output_file: str):
    """
    Основная функция для создания документа Word.
    
    :param formatted_results: Отформатированная строка с результатами
    :param output_file: Путь для сохранения документа Word
    """
    logging.info("Начало создания документа Word")
    create_word_document(formatted_results, output_file)
    logging.info("Создание документа Word завершено")

