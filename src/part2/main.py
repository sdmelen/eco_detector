import pandas as pd
from docx import Document
from datetime import datetime, timedelta


def get_duration_string(total_minutes):
    """Возвращает строку длительности в часах и минутах."""
    hours = total_minutes // 60
    minutes = total_minutes % 60
    hours_str = f"{hours} часов" if hours > 0 else ""
    minutes_str = f"{minutes} минут" if minutes > 0 else ""
    return f"{hours_str} {minutes_str}".strip()


def adjust_datetime(date_time_str, minutes_to_subtract):
    """Вычитает заданное количество минут из строки даты и времени."""
    date_time_format = "%d/%m/%Y %H:%M"
    original_date_time = datetime.strptime(date_time_str, date_time_format)
    new_date_time = original_date_time - timedelta(minutes=minutes_to_subtract)
    return new_date_time.strftime(date_time_format)


def add_paragraph_to_document(document, file_name, duration_str, details):
    """Добавляет абзац в документ с заданными параметрами."""
    paragraph_text = f"по {file_name} - {duration_str} {details}"
    document.add_paragraph(paragraph_text)


def record_max_excess_duration(file_name, document):
    """Находит строки с максимальным количеством точек и записывает результаты в документ."""
    df = pd.read_excel(f"{file_name}.xlsx")
    max_points = df['Количество точек'].max()
    max_excess_rows = df[df['Количество точек'] == max_points].values.tolist()

    for line in max_excess_rows:
        duration_minutes = int(line[2]) * 20
        duration_str = get_duration_string(duration_minutes)
        adjusted_start_time = adjust_datetime(line[3], 20)
        details = f"с {adjusted_start_time} по {line[4]} ({line[0]})"
        add_paragraph_to_document(document, file_name, duration_str, details)

    document.save('test.docx')


def record_total_excess_duration(file_name, document):
    """Находит суммарную длительность превышений для каждой точки и записывает результаты в документ."""
    df = pd.read_excel(f"{file_name}.xlsx")
    points_map = df.groupby(df.columns[0])['Количество точек'].sum().to_dict()
    max_key = max(points_map, key=points_map.get)

    total_duration_minutes = points_map[max_key] * 20
    duration_str = get_duration_string(total_duration_minutes)
    details = f"({max_key})"
    add_paragraph_to_document(document, file_name, duration_str, details)

    document.save('test.docx')


def main():
    """Основная функция для выполнения всех операций."""
    document = Document()
    document.add_heading('Максимальная непрерывная длительность превышений:', level=0)

    file_name = 'H2S_период'
    record_max_excess_duration(file_name, document)

    document.add_heading('Максимальная общая длительность превышений:', level=0)
    record_total_excess_duration(file_name, document)

    for paragraph in document.paragraphs:
        print(paragraph.text)


if __name__ == "__main__":
    main()
